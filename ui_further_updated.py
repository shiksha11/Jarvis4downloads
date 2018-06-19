import sys
from tkinter import *
import os
import json
import codecs
import PyPDF2
import docx
import csv
from xml.etree import ElementTree
import speech_recognition as sr
import pyaudio



root = Tk()
root.title("Search Box")
root.geometry("1150x900+750+250")
root.configure(background="#00FFFF")


cwd = os.getcwd()
o=cwd.split(os.sep)
os.chdir(os.sep+ o[1] + os.sep + o[2] + os.sep)
p=os.getcwd()
t=os.sep
json_file =open(os.sep+ o[1] + os.sep + o[2] + os.sep +'config.txt',"r")
data =  json.load(json_file)
json_file.close()


def filename():
    global path

    button6= Button(root, text= "Search by filename" ,fg= "red", command = calling_function1,bg="#00FFFF",font="Times 18 bold")
    button6.pack()
    button7= Button(root, text ="Body Search" ,fg= "red", command = calling_function2,bg="#00FFFF",font="Times 18 bold")
    button7.pack()


def calling_function1():
    global path
    p=entry1.get()
    Search(p)
def calling_function2():
    global path
    p=entry1.get()
    documents2()
def calling_function3():
    global path
    global file_input
    Search(file_input)
def calling_function4():
    global path
    global file_input
    documents3(file_input)


def Search(p):
    global file
    global path
    global found
    found=False
    file = p
    f_name, f_ext = os.path.splitext(file)

    f_name=f_name.lower()
    if(f_ext !=''):
        if(any( x== f_ext for x in data['pictures'])):
            for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'pictures'):
                for p in files:
                    temp=p
                    file=file.lower()
                    p=p.lower()
                    if(p==file):
                        found=True
                        path=(os.path.join(root,temp))
                        printing(path)
            if(found==False):
                wordList=re.sub("[^\w]", " ", f_name).split()
                for i in wordList:
                    if('_' in i):
                        j=i.split('_')
                        wordList.remove(i)
                        for t in j:
                            wordList.append(t)
                d={}
                num=[]
                for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'pictures'):
                    for p in files:
                        temp=p
                        p=p.lower()
                        fileList=re.sub("[^\w]", " ", p).split()
                        for i in fileList:
                            if('_' in i):
                                j=i.split('_')
                                fileList.remove(i)
                                for t in j:
                                    fileList.append(t)
                        count=0
                        for i in wordList:
                            if(any(i==j for j in fileList)):
                                count+=1
                        d[temp]=count
                common_words=[]
                for i in d.values():
                    common_words.append(i)
                common_words=sorted(common_words)

                for i in range(max(d.values()),0,-1):
                    for key,value in d.items():
                        if(value==i):
                            found=True
                            path=os.path.join(root,key)
                            printing(path)
            if(found==False):
               d={}
               count=[]
               for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'pictures'):
                   for p in files:
                       temp=p
                       p=p.lower()
                       list1=set([i for i in p])
                       list2=set([i for i in f_name])
                       f_name=f_name.lower()
                       new = set(p+f_name)
                       new1 = set(list1.intersection(list2))
                       f = (len(new)-len(new1))/len(new)
                       count.append(f)
                       d[temp]=f
               count = sorted(count)
               t=count[0]
               for key,value in d.items():
                   if(value==t):
                       found = True
                       path=os.path.join(root,key)
                       print(os.path.join(root,key))
                       printing(path)

        elif(any( x == f_ext for x in data['videos'])):
            for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'videos'):
                for p in files:
                    temp=p
                    file=file.lower()
                    p=p.lower()
                    if(p==file):
                        found=True
                        path=(os.path.join(root,temp))
                        printing(path)
            if(found==False):
                wordList=re.sub("[^\w]", " ", f_name).split()
                for i in wordList:
                    if('_' in i):
                        j=i.split('_')
                        wordList.remove(i)
                        for t in j:
                            wordList.append(t)
                d={}
                num=[]
                for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'videos'):
                    for p in files:
                        temp=p
                        p=p.lower()
                        fileList=re.sub("[^\w]", " ", p).split()
                        for i in fileList:
                            if('_' in i):
                                j=i.split('_')
                                fileList.remove(i)
                                for t in j:
                                    fileList.append(t)
                        count=0
                        for i in wordList:
                            if(any(i==j for j in fileList)):
                                count+=1
                        d[temp]=count
                common_words=[]
                for i in d.values():
                    common_words.append(i)
                common_words=sorted(common_words)

                for i in range(max(d.values()),0,-1):
                    for key,value in d.items():
                        if(value==i):
                            found=True

                            path=os.path.join(root,key)

                            printing(path)
            if(found==False):
               d={}
               count=[]
               for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'videos'):
                   for p in files:
                       temp=p
                       p=p.lower()
                       list1=set([i for i in p])
                       list2=set([i for i in f_name])
                       f_name=f_name.lower()
                       new = set(p+f_name)
                       new1 = set(list1.intersection(list2))
                       f = (len(new)-len(new1))/len(new)
                       count.append(f)
                       d[temp]=f
               count = sorted(count)
               t=count[0]
               for key,value in d.items():
                   if(value==t):
                       found = True
                       path=os.path.join(root,key)
                       print(os.path.join(root,key))
                       printing(path)




        elif(any( x== f_ext for x in data['documents'])):
            for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'documents'):
                for p in files:
                    temp=p
                    file=file.lower()
                    p=p.lower()
                    if(p==file):
                        found=True
                        path=(os.path.join(root,temp))
                        print(path)
                        printing(path)
            if(found==False):
                wordList=re.sub("[^\w]", " ", f_name).split()
                for i in wordList:
                    if('_' in i):
                        j=i.split('_')
                        wordList.remove(i)
                        for t in j:
                            wordList.append(t)
                d={}
                num=[]
                for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'documents'):
                    for p in files:
                        temp=p
                        p=p.lower()
                        fileList=re.sub("[^\w]", " ", p).split()
                        for i in fileList:
                            if('_' in i):
                                j=i.split('_')
                                fileList.remove(i)
                                for t in j:
                                    fileList.append(t)
                        count=0
                        for i in wordList:
                            if(any(i==j for j in fileList)):
                                count+=1
                        d[temp]=count
                common_words=[]
                for i in d.values():
                    common_words.append(i)
                common_words=sorted(common_words)

                for i in range(max(d.values()),0,-1):
                    for key,value in d.items():
                        if(value==i):
                            found = True
                            path=os.path.join(root,key)
                            print(os.path.join(root,key))
                            printing(path)
                if(found==False):
                   d={}
                   count=[]
                   for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'documents'):
                       for p in files:
                           temp=p
                           p=p.lower()
                           list1=set([i for i in p])
                           list2=set([i for i in f_name])
                           f_name=f_name.lower()
                           new = set(p+f_name)
                           new1 = set(list1.intersection(list2))
                           f = (len(new)-len(new1))/len(new)
                           count.append(f)
                           d[temp]=f
                   count = sorted(count)
                   t=count[0]
                   for key,value in d.items():
                       if(value==t):
                           found = True
                           path=os.path.join(root,key)
                           print(os.path.join(root,key))
                           printing(path)




        else:
            for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'projects'):
                for p in files:
                    temp=p
                    file=file.lower()
                    p=p.lower()
                    if(p==file):
                        found=True
                        path=(os.path.join(root,temp))
                        printing(path)
            if(found==False):
                wordList=re.sub("[^\w]", " ", f_name).split()
                for i in wordList:
                    if('_' in i):
                        j=i.split('_')
                        wordList.remove(i)
                        for t in j:
                            wordList.append(t)
                d={}
                num=[]
                for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'projects'):
                    for p in files:
                        temp=p
                        p=p.lower()
                        fileList=re.sub("[^\w]", " ", p).split()
                        for i in fileList:
                            if('_' in i):
                                j=i.split('_')
                                fileList.remove(i)
                                for t in j:
                                    fileList.append(t)
                        count=0
                        for i in wordList:
                            if(any(i==j for j in fileList)):
                                count+=1
                        d[temp]=count
                common_words=[]
                for i in d.values():
                    common_words.append(i)
                common_words=sorted(common_words)

                for i in range(max(d.values()),0,-1):
                    for key,value in d.items():
                        if(value==i):
                            found = True
                            path=os.path.join(root,key)
                            printing(path)
            if(found==False):
                d={}
                count=[]
                for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'projects'):
                    for p in files:
                        temp=p
                        p=p.lower()
                        list1=set([i for i in p])
                        list2=set([i for i in f_name])
                        f_name=f_name.lower()
                        new = set(p+f_name)
                        new1 = set(list1.intersection(list2))
                        f = (len(new)-len(new1))/len(new)
                        count.append(f)
                        d[temp]=f
                count = sorted(count)
                t=count[0]
                for key,value in d.items():
                    if(value==t):
                        found = True
                        path=os.path.join(root,key)
                        print(os.path.join(root,key))
                        printing(path)
    if(f_ext==''):
        printing("File name has no extension.")
        printing("Select the file type.")

def printing(p):
    thelabel= Label(root, text=p,bg="#00FFFF",font="none 16 bold").pack()
def pictures():
    global f_name
    f_name, f_ext = os.path.splitext(file)

    wordList=re.sub("[^\w]", " ", f_name).split()
    for i in wordList:
        if('_' in i):
            j=i.split('_')
            wordList.remove(i)
            for t in j:
                wordList.append(t)
    d={}
    num=[]
    for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'pictures'):
        for p in files:
            temp=p
            p=p.lower()
            fileList=re.sub("[^\w]", " ", p).split()
            for i in fileList:
                if('_' in i):
                    j=i.split('_')
                    fileList.remove(i)
                    for t in j:
                        fileList.append(t)
            count=0
            for i in wordList:
                if(any(i==j for j in fileList)):
                    count+=1
            d[temp]=count
    common_words=[]
    for i in d.values():
        common_words.append(i)
    common_words=sorted(common_words)

    for i in range(max(d.values()),0,-1):
        for key,value in d.items():
            if(value==i):
                found = True
                path=os.path.join(root,key)
                printing(path)
    if(found==False):
        d={}
        count=[]
        for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'pictures'):
            for p in files:
                temp=p
                p=p.lower()
                list1=set([i for i in p])
                list2=set([i for i in f_name])
                f_name=f_name.lower()
                new = set(p+f_name)
                new1 = set(list1.intersection(list2))
                f = (len(new)-len(new1))/len(new)
                count.append(f)
                d[temp]=f
        count = sorted(count)
        t=count[0]
        for key,value in d.items():
            if(value==t):
                found = True
                path=os.path.join(root,key)
                print(os.path.join(root,key))
                printing(path)







def videos():
    f_name, f_ext = os.path.splitext(file)
    global found
    wordList=re.sub("[^\w]", " ", f_name).split()
    for i in wordList:
        if('_' in i):
            j=i.split('_')
            wordList.remove(i)
            for t in j:
                wordList.append(t)
    d={}
    num=[]
    for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'videos'):
        for p in files:
            temp=p
            p=p.lower()
            fileList=re.sub("[^\w]", " ", p).split()
            for i in fileList:
                if('_' in i):
                    j=i.split('_')
                    fileList.remove(i)
                    for t in j:
                        fileList.append(t)
            count=0
            for i in wordList:
                if(any(i==j for j in fileList)):
                    count+=1
            d[temp]=count
    common_words=[]
    for i in d.values():
        common_words.append(i)
    common_words=sorted(common_words)

    for i in range(max(d.values()),0,-1):
        for key,value in d.items():
            if(value==i):
                found = True
                path=os.path.join(root,key)
                printing(path)
    if(found==False):
        d={}
        count=[]
        for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'videos'):
            for p in files:
                temp=p
                p=p.lower()
                list1=set([i for i in p])
                list2=set([i for i in f_name])
                f_name=f_name.lower()
                new = set(p+f_name)
                new1 = set(list1.intersection(list2))
                f = (len(new)-len(new1))/len(new)
                count.append(f)
                d[temp]=f
        count = sorted(count)
        t=count[0]
        for key,value in d.items():
            if(value==t):
                found = True
                path=os.path.join(root,key)
                print(os.path.join(root,key))
                printing(path)
def documents():
    f_name, f_ext = os.path.splitext(file)
    global found
    wordList=re.sub("[^\w]", " ", f_name).split()
    for i in wordList:
        if('_' in i):
            j=i.split('_')
            wordList.remove(i)
            for t in j:
                wordList.append(t)
    d={}
    num=[]
    for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'documents'):
        for p in files:
            temp=p
            p=p.lower()
            fileList=re.sub("[^\w]", " ", p).split()
            for i in fileList:
                if('_' in i):
                    j=i.split('_')
                    fileList.remove(i)
                    for t in j:
                        fileList.append(t)
            count=0
            for i in wordList:
                if(any(i==j for j in fileList)):
                    count+=1
            d[temp]=count
    common_words=[]
    for i in d.values():
        common_words.append(i)
    common_words=sorted(common_words)

    for i in range(max(d.values()),0,-1):
        for key,value in d.items():
            if(value==i):
                found = True
                path=os.path.join(root,key)
                printing(path)
    if(found==False):
        d={}
        count=[]
        for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'documents'):
            for p in files:
                temp=p
                p=p.lower()
                list1=set([i for i in p])
                list2=set([i for i in f_name])
                f_name=f_name.lower()
                new = set(p+f_name)
                new1 = set(list1.intersection(list2))
                f = (len(new)-len(new1))/len(new)
                count.append(f)
                d[temp]=f
        count = sorted(count)
        t=count[0]
        for key,value in d.items():
            if(value==t):
                found = True
                path=os.path.join(root,key)
                print(os.path.join(root,key))
                printing(path)
        if(found==False):
            d={}
            count=[]
            for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'documents'):
                for p in files:
                    temp=p
                    p=p.lower()
                    list1=set([i for i in p])
                    list2=set([i for i in f_name])
                    f_name=f_name.lower()
                    new = set(p+f_name)
                    new1 = set(list1.intersection(list2))
                    f = (len(new)-len(new1))/len(new)
                    count.append(f)
                    d[temp]=f
            count = sorted(count)
            t=count[0]
            for key,value in d.items():
                if(value==t):
                    found = True
                    path=os.path.join(root,key)
                    print(os.path.join(root,key))
                    printing(path)

def projects():
    global found
    global f_name
    found = False
    f_name, f_ext = os.path.splitext(file)
    wordList=re.sub("[^\w]", " ", f_name).split()
    for i in wordList:
        if('_' in i):
            j=i.split('_')
            wordList.remove(i)
            for t in j:
                wordList.append(t)
    d={}
    num=[]
    for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'projects'):
        for p in files:
            temp=p
            p=p.lower()
            fileList=re.sub("[^\w]", " ", p).split()
            for i in fileList:
                if('_' in i):
                    j=i.split('_')
                    fileList.remove(i)
                    for t in j:
                        fileList.append(t)
            count=0
            for i in wordList:
                if(any(i==j for j in fileList)):
                    count+=1
            d[temp]=count
    common_words=[]
    for i in d.values():
        common_words.append(i)
    common_words=sorted(common_words)

    for i in range(max(d.values()),0,-1):
        for key,value in d.items():
            if(value==i):
                found = True
                path=os.path.join(root,key)
                printing(path)
    if(found==False):
        d={}
        count=[]
        for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'projects'):
            for p in files:

                temp=p
                p=p.lower()
                list1=set([i for i in p])
                list2=set([i for i in f_name])
                f_name=f_name.lower()
                new = set(p+f_name)
                new1 = set(list1.intersection(list2))
                f = (len(new)-len(new1))/len(new)
                count.append(f)
                d[temp]=f
        count = sorted(count)
        t=count[0]
        for key,value in d.items():
            if(value==t):
                found = True
                path=os.path.join(root,key)
                print(os.path.join(root,key))
                printing(path)


def documents2():
    q=entry1.get()
    wordList=re.sub("[^\w]", " ", q).split()

    d={}
    counts = []
    for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'documents'):
        for p in files:

            count = 0
            y= os.path.join(root,p)
            f_name, f_ext = os.path.splitext(y)
            if(f_ext =='.txt'):
                with codecs.open(y, "r" , encoding ="utf-8", errors = 'ignore') as file1:
                    for line in file1:
                        line = line.rstrip()
                        words = line.split()
                        for w in words:
                            w=w.lower()
                            if(any(x.lower() == w for x in wordList)):
                                count+=1
                print(count)
                counts.append(count)
                print(count)
                d[p] = count
            elif(f_ext=='.pdf'):
                file = open(os.sep + o[1] + os.sep + o[2]+ os.sep +'Desktop'+ os.sep+'documents'+ os.sep + p , 'rb')
                pdfreader = PyPDF2.PdfFileReader(file)
                total = pdfreader.getNumPages()
                for i in range(total):
                    pageobj = pdfreader.getPage(i)
                    text = pageobj.extractText()
                    words = re.sub("[^\w]", " ", text).split()

                    for i in words:
                        i = i.lower()

                        if(any(i==x.lower() for x in wordList)):
                            count+=1
                print(count)
                counts.append(count)
                d[p] = count
            elif(f_ext=='.docx'):
                document = docx.Document(os.sep + o[1] + os.sep + o[2] + os.sep +'Desktop'+ os.sep+'documents'+ os.sep +p)
                for i in document.paragraphs:
                    v=i.text
                    words = re.sub("[^\w]", " ", v).split()

                    for i in words:
                        i = i.lower()

                        if(any(i==x.lower() for x in wordList)):
                            count+=1
                print(count)
                counts.append(count)
                d[p] = count
            elif(f_ext == '.csv'):
                file_path = os.sep + o[1] + os.sep + o[2]+ os.sep +'Desktop'+ os.sep+'documents'+ os.sep + p
                file = open(file_path, newline='')
                reader = csv.reader(file)
                data = [row for row in reader]

                for i in range(len(data)):
                    r = data[i]
                    for j in range(len(data[i])):
                        m=data[i][j]
                        if(any(m.lower() == x for x in wordList)):
                            count+=1
                print(count)
                counts.append(count)
                d[p] = count


    t=max(counts)

    for key,value in d.items():
        if(value==t):
            found = True
            path=os.path.join(root,key)
            print(os.path.join(root,key))
            printing(path)
def documents3(file_input):
    q=file_input
    wordList=re.sub("[^\w]", " ", file_input).split()

    d={}
    counts = []
    for root,dirs,files in os.walk(os.sep + o[1] + os.sep + o[2] + os.sep + 'Desktop' + os.sep+ 'documents'):
        for p in files:

            count = 0
            y= os.path.join(root,p)
            f_name, f_ext = os.path.splitext(y)
            if(f_ext =='.txt'):
                with codecs.open(y, "r" , encoding ="utf-8", errors = 'ignore') as file1:
                    for line in file1:
                        line = line.rstrip()
                        words = line.split()
                        for w in words:
                            w=w.lower()
                            if(any(x.lower() == w for x in wordList)):
                                count+=1
                print(count)
                counts.append(count)
                print(count)
                d[p] = count
            elif(f_ext=='.pdf'):
                file = open(os.sep + o[1] + os.sep + o[2]+ os.sep +'Desktop'+ os.sep+'documents'+ os.sep + p , 'rb')
                pdfreader = PyPDF2.PdfFileReader(file)
                total = pdfreader.getNumPages()
                for i in range(total):
                    pageobj = pdfreader.getPage(i)
                    text = pageobj.extractText()
                    words = re.sub("[^\w]", " ", text).split()

                    for i in words:
                        i = i.lower()

                        if(any(i==x.lower() for x in wordList)):
                            count+=1
                print(count)
                counts.append(count)
                d[p] = count
            elif(f_ext=='.docx'):
                document = docx.Document(os.sep + o[1] + os.sep + o[2] + os.sep +'Desktop'+ os.sep+'documents'+ os.sep +p)
                for i in document.paragraphs:
                    v=i.text
                    words = re.sub("[^\w]", " ", v).split()

                    for i in words:
                        i = i.lower()

                        if(any(i==x.lower() for x in wordList)):
                            count+=1
                print(count)
                counts.append(count)
                d[p] = count
            elif(f_ext == '.csv'):
                file_path = os.sep + o[1] + os.sep + o[2]+ os.sep +'Desktop'+ os.sep+'documents'+ os.sep + p
                file = open(file_path, newline='')
                reader = csv.reader(file)
                data = [row for row in reader]

                for i in range(len(data)):
                    r = data[i]
                    for j in range(len(data[i])):
                        m=data[i][j]
                        if(any(m.lower() == x for x in wordList)):
                            count+=1
                print(count)
                counts.append(count)
                d[p] = count


    t=max(counts)

    for key,value in d.items():
        if(value==t):
            found = True
            path=os.path.join(root,key)
            print(os.path.join(root,key))
            printing(path)

def speech():
    global file_input
    thelabel7 = Label(root, text="Listening..", fg="black",bg="#00FFFF", font="Times 18 bold underline")
    thelabel7.pack()
    r= sr.Recognizer()
    with sr.Microphone() as source:

        audio = r.listen(source)
        file_input = r.recognize_google(audio)
        thelabel8 = Label(root, text = file_input, fg="black",bg="#00FFFF", font="Times 18 bold underline")
        thelabel8.pack()
        button11= Button(root, text= "Search by filename" ,fg= "red", command = calling_function3,bg="#00FFFF",font="Times 18 bold")
        button11.pack()
        button12= Button(root, text ="Body Search" ,fg= "red", command = calling_function4,bg="#00FFFF",font="Times 18 bold")
        button12.pack()









thelabel2 = Label(root, text="Enter File Name:", fg="black",bg="#00FFFF", font="Times 18 bold underline")
thelabel2.pack()
entry1= Entry(root)
entry1.pack()

button1= Button(root, text ="Search Path.." ,fg= "red", command = filename,bg="#00FFFF",font="Times 18 bold")
button1.pack()
button2= Button(root, text ="Pictures" ,fg= "red", command = pictures,bg="#00FFFF",font="Times 18 bold")
button2.pack(side=LEFT)
button3= Button(root, text ="  Videos" ,fg= "red", command = videos,bg="#00FFFF",font="Times 18 bold")
button3.pack(side=LEFT)
button4= Button(root, text ="Projects" ,fg= "red", command = projects,bg="#00FFFF",font="Times 18 bold")
button4.pack(side=RIGHT)
button5= Button(root, text ="Documents" ,fg= "red", command = documents,bg="#00FFFF",font="Times 18 bold")
button5.pack(side=RIGHT)

button5= Button(root, text ="Give Input in Audio Form." ,fg= "red", command = speech,bg="#00FFFF",font="Times 18 bold")
button5.pack(side=RIGHT)







root.mainloop()
